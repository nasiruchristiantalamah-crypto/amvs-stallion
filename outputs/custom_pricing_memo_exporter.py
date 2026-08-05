"""
================================================================================
CUSTOM PRODUCT PRICING — full Actuarial Memorandum (Word)
================================================================================
What this file does:
    Writes the "Generate Memorandum" button's Word document for the
    dashboard's Pricing page — a full actuarial pricing memorandum, in the
    same structure as a real signed-off memo (Policy Overview, Premium
    Estimation basis, Appendix rate table, Illustrative Example, Actuarial
    Declaration; see tests/test_pricing_real_world.py for the real Impact
    Life Afentoboa Plus memo this structure is modelled on).

    Two kinds of content go into it:
      - Narrative sections (Policy Overview, Benefits notes, Proceeds/
        Settlement, Underwriting, Lapsation/Reinstatement, Closing notes) —
        supplied by the user via the dashboard's text areas, since AVMS has
        no way to invent contractual/legal product language on its own.
        Each has a plain auto-composed fallback if left blank, so the
        button always produces something usable.
      - Computed sections (Policy Benefits table, Premium Estimation basis,
        Appendix rate table, Illustrative Example breakdown) — built
        entirely from the SAME engine output every other export on this
        page uses (engine.custom_pricing.run_custom_pricing /
        run_custom_rate_table), so the numbers in this memo can never
        drift from what the dashboard actually priced.

    This is a heavier, more complete document than
    outputs/custom_pricing_word_exporter.py's short AVR product note —
    that one stays as the quick single-page summary; this one is the full
    memorandum.
================================================================================
"""

import os
from datetime import datetime
from typing import Dict, List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from engine.assumptions import ProductAssumptions
from engine.product import Product
from engine.decrement import run_decrement_projection
from engine.cashflows import calculate_cash_flows

GENERATED_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "outputs", "generated")


def _heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def _para(doc: Document, text: str, italic: bool = False, bold: bool = False, size: Optional[int] = None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.italic = italic
    run.bold = bold
    if size:
        run.font.size = Pt(size)


def _kv_table(doc: Document, rows: List[tuple]) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.style = "Light Grid Accent 1"
    for k, v in rows:
        row = table.add_row().cells
        row[0].text = str(k)
        row[1].text = str(v)


def _fmt_pct(x) -> str:
    try:
        return f"{float(x) * 100:.2f}%"
    except (TypeError, ValueError):
        return str(x)


def _fmt_money(x) -> str:
    try:
        return f"GHS {float(x):,.2f}"
    except (TypeError, ValueError):
        return str(x)


def _month_one_breakdown(product: Product, assumptions: ProductAssumptions, monthly_premium: float) -> dict:
    """
    Re-derive the Pure Risk Premium / Policy Fee / Expenses / Commission /
    Profit Margin split at the product's own entry age, mirroring a real
    memo's "Illustrative Example" table. Month 1 has lx=1 for the main life
    by construction (nobody has died or lapsed yet), so its CashFlowRow
    fields ARE the per-policy figures needed directly — no new engine
    return-shape changes required.
    """
    dec_rows = run_decrement_projection(assumptions, product)
    cf_rows = calculate_cash_flows(dec_rows, assumptions, product, monthly_premium)
    month_one = cf_rows[0]
    policy_fee = assumptions.policy_fee_monthly
    expenses_excl_fee = max(0.0, month_one.total_expenses - policy_fee)
    profit = month_one.gross_premium - month_one.total_benefits - month_one.total_expenses - month_one.commission
    return {
        "pure_risk_premium": month_one.total_benefits,
        "policy_fee": policy_fee,
        "expenses": expenses_excl_fee,
        "commission": month_one.commission,
        "profit_margin": profit,
        "total": month_one.gross_premium,
    }


def _auto_policy_overview(product_spec: dict) -> str:
    name = product_spec.get("product_name", "this product")
    term = product_spec.get("policy_term_years")
    term_txt = f"a {term}-year renewable term" if term else "a whole-of-life"
    riders = product_spec.get("riders") or []
    rider_names = ", ".join(r.get("name", "") for r in riders) if riders else "no additional riders"
    dep_count = len(product_spec.get("dependants") or [])
    return (
        f"{name} is {term_txt} insurance product. The plan provides a main sum assured "
        f"of {_fmt_money(product_spec.get('sum_assured', 0))}, plus {rider_names}. "
        f"The plan covers the main life and up to {dep_count} additional insured "
        f"{'life' if dep_count == 1 else 'lives'}."
    )


def generate_actuarial_memorandum(
    result:             dict,
    rate_table:         Dict[int, dict],
    product_spec:       dict,
    narrative:          Optional[dict] = None,
    appointed_actuary:  str = "Charles Osei-Akoto, ASA, MAAA",
    consulting_firm:    str = "Stallion Consultants Ltd",
    report_date:        Optional[str] = None,
    product:            Optional[Product] = None,
    assumptions:        Optional[ProductAssumptions] = None,
    output_path:        Optional[str] = None,
) -> str:
    """
    Build a full actuarial pricing memorandum and write it to disk.

    Returns:
        The path the document was written to.
    """
    narrative = narrative or {}
    report_date = report_date or datetime.now().strftime("%d %B %Y")
    a = result.get("assumptions_used", {})
    product_name = result.get("product_name") or product_spec.get("product_name", "Custom Product")

    doc = Document()
    section = doc.sections[0]
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    # ── Cover page ──────────────────────────────────────────────────────
    title = doc.add_heading("ACTUARIAL MEMORANDUM", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run(f"FOR {product_name.upper()}")
    sub_run.bold = True
    sub_run.font.size = Pt(14)
    _para(doc, report_date, italic=True)
    doc.add_paragraph()
    _para(doc, f"Presented By:\n{consulting_firm}", bold=True)
    doc.add_page_break()

    # ── I. Policy Overview ───────────────────────────────────────────────
    _heading(doc, "I. Policy Overview")
    _para(doc, narrative.get("policy_overview") or _auto_policy_overview(product_spec))

    _heading(doc, "A. Number of Insureds", level=2)
    dependants = product_spec.get("dependants") or []
    _para(doc, f"This policy covers the Main Life{' as well as ' + str(len(dependants)) + ' additional insured life(ves)' if dependants else ' only'}.")

    _heading(doc, "B. Policy Benefits", level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = "Benefit", "Basis", "Amount (GHS)"
    row = table.add_row().cells
    row[0].text, row[1].text, row[2].text = "Main Benefit (Sum Assured)", product_spec.get("product_type", ""), str(product_spec.get("sum_assured", 0))
    for r in (product_spec.get("riders") or []):
        row = table.add_row().cells
        row[0].text = str(r.get("name", ""))
        row[1].text = str(r.get("benefit_type", ""))
        row[2].text = str(r.get("benefit_amount", ""))

    _heading(doc, "C-E. Policy Proceeds, Surrender & Settlement", level=2)
    _para(doc, narrative.get("proceeds_and_settlement") or
          "Proceeds are payable as a lump sum upon the occurrence of an insured event, equal to the "
          "applicable benefit amount(s) less any outstanding premium required to maintain cover to the date "
          "of the event.")

    _heading(doc, "F. Premium", level=2)
    _para(doc, narrative.get("premium_notes") or
          f"Premiums are payable {product_spec.get('premium_mode', 'monthly')}. The premium covers the risk "
          f"cost of the benefits under this policy, together with policy fees, expenses, commission, and "
          f"the target profit margin.")

    _heading(doc, "G. Underwriting Conditions", level=2)
    _para(doc, narrative.get("underwriting_notes") or
          f"Entry age: {product_spec.get('entry_age', '—')}. Gender basis: {product_spec.get('gender', 'unisex')}. "
          f"Policy term: {product_spec.get('policy_term_years') or 'Whole of life'}.")

    _heading(doc, "H-I. Lapsation & Reinstatement", level=2)
    _para(doc, narrative.get("lapsation_and_reinstatement") or
          "If premiums are not paid, the policy is subject to a grace period before lapsing. A lapsed policy "
          "may be eligible for reinstatement subject to the insurer's terms and payment of arrears.")

    doc.add_page_break()

    # ── II. Premium Estimation ───────────────────────────────────────────
    _heading(doc, "II. Premium Estimation")
    _para(doc, "The premium was determined using the equivalence principle: expected cash flows "
               "(premiums, benefits, commission, expenses) were projected and discounted, and the "
               "premium solved for that achieves the target profit margin.")

    _heading(doc, "A. Insurance Risk Assumptions", level=2)
    _kv_table(doc, [
        ("Mortality basis", a.get("mortality_table", "SA 85/90 (Actuarial Society of South Africa)")),
        ("Mortality loading", _fmt_pct(a.get("mortality_loading"))),
        ("Gender basis", a.get("gender_main_str")),
        ("Valuation interest rate (p.a.)", _fmt_pct(a.get("valuation_rate_pa"))),
        ("Investment rate of return (p.a.)", _fmt_pct(a.get("investment_rate_pa"))),
        ("Expense inflation rate (p.a.)", _fmt_pct(a.get("expense_inflation_pa"))),
        ("Collection rate", _fmt_pct(a.get("collection_rate"))),
        ("Acquisition cost (GHS)", a.get("acquisition_cost")),
        ("Renewal expense (GHS p.a.)", a.get("renewal_expense_annual")),
        ("Policy fee (GHS/month)", a.get("policy_fee_monthly")),
        ("Target profit margin", _fmt_pct(a.get("target_profit_margin"))),
    ])
    commission = a.get("commission") or {}
    _kv_table(doc, [
        ("Commission — initial year", _fmt_pct(commission.get("initial_rate"))),
        ("Commission — renewal years", _fmt_pct(commission.get("renewal_rate"))),
    ])
    lapse = (a.get("lapse_schedule") or {}).get("rates", {})
    if lapse:
        _para(doc, "Lapse rates:")
        table = doc.add_table(rows=1, cols=2)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text, hdr[1].text = "Policy year", "Lapse rate"
        for yr in sorted(lapse.keys(), key=lambda x: int(x)):
            row = table.add_row().cells
            row[0].text = str(yr)
            row[1].text = _fmt_pct(lapse[yr])

    doc.add_page_break()

    # ── III. Appendix: Monthly Risk Premium ──────────────────────────────
    _heading(doc, "III. Appendix: Monthly Risk Premium")
    _para(doc, "The monthly risk premium by issue age of the main life is shown below.")
    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = "Age", "Monthly Premium (GHS)", "Onerous?"
    for age in sorted(rate_table.keys()):
        entry = rate_table[age]
        row = table.add_row().cells
        row[0].text = str(age)
        row[1].text = str(entry.get("error") or entry.get("monthly_premium"))
        row[2].text = "Yes" if entry.get("is_onerous") else ("—" if entry.get("error") else "No")

    doc.add_page_break()

    # ── IV. Illustrative Example ──────────────────────────────────────────
    _heading(doc, "IV. Illustrative Example")
    if product is not None and assumptions is not None:
        breakdown = _month_one_breakdown(product, assumptions, result.get("monthly_premium", 0))
        _para(doc, f"Breakdown of the monthly premium of {_fmt_money(result.get('monthly_premium'))} "
                   f"at entry age {product_spec.get('entry_age')}:")
        table = doc.add_table(rows=0, cols=2)
        table.style = "Light Grid Accent 1"
        for label, key in [
            ("Pure Risk Premium", "pure_risk_premium"), ("Policy fee", "policy_fee"),
            ("Expenses", "expenses"), ("Commission", "commission"),
            ("Profit Margin", "profit_margin"), ("Monthly Premium (Total)", "total"),
        ]:
            row = table.add_row().cells
            row[0].text = label
            row[1].text = _fmt_money(breakdown[key])

    _heading(doc, "Additional Notes", level=2)
    _para(doc, narrative.get("closing_notes") or
          "Actual experience (mortality, morbidity, lapse, expenses, and investment returns) may differ "
          "from the assumptions above; such variances will emerge as additional profit or loss over the "
          "life of the policy.")

    doc.add_page_break()

    # ── V. Actuarial Declaration ──────────────────────────────────────────
    _heading(doc, "V. Actuarial Declaration")
    _para(doc, "This is to certify that the methodology used in pricing this policy is in accordance with "
               "internationally accepted actuarial principles.")
    doc.add_paragraph()
    doc.add_paragraph("-" * 30 + "\t\t" + report_date)
    p = doc.add_paragraph()
    run = p.add_run(appointed_actuary)
    run.bold = True
    _para(doc, "(Consulting Actuary)", italic=True)
    _para(doc, consulting_firm)

    if output_path is None:
        os.makedirs(GENERATED_DIR, exist_ok=True)
        name_slug = "".join(c if c.isalnum() else "_" for c in product_name)[:40]
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_path = os.path.join(GENERATED_DIR, f"ActuarialMemo_{name_slug}_{timestamp}.docx")

    doc.save(output_path)
    return output_path
