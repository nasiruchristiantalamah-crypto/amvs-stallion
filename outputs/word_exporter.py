"""
================================================================================
NIC AVR REPORT — WORD (.docx) EXPORTER
================================================================================
What this file does:
    Renders the complete NIC Actuarial Valuation Report (AVR) as a Word
    document, using python-docx. Two data sources feed it:

        1. api.nic_report.generate_avr_data() — the life-side 9-section
           JSON model (executive summary, product description, assumptions,
           methodology, LRC/LIC, income statement, CSM roll-forward,
           solvency, actuarial opinion).
        2. engine.ifrs17_nonlife.generate_nonlife_paa_statements() +
           engine.journals.generate_nonlife_journal() directly — RICHER
           non-life content than the older avr["section5_non_life"] (which
           only carries the 4-metric IBNR/OCR/ULAE/UPR/DAC table from
           run_nic_summary()). This module instead builds the same
           Balance Sheet / Income Statement / journal listing that
           outputs/excel_exporter.py produces, so the Word and Excel
           non-life exports show the same numbers in the same shape.

    Applies Stallion Consultants branding (navy headings matching
    outputs/excel_exporter.py's STALLION_NAVY, professional gridded
    tables with a navy header row) and ends with an actuarial certificate
    page carrying a signature block, mirroring section 9's opinion text.

Honest scope note:
    This assembles and formats numbers the ENGINE has already computed and
    validated elsewhere (see engine/ifrs17.py, engine/ifrs17_nonlife.py,
    and their test suites) — it does no calculation of its own. Any
    modelling caveat documented in those modules (e.g. the RA/discounting
    simplifications in engine/ifrs17_nonlife.py, or a client's non-life
    data not yet being benchmark-validated — see clients/qic/client.yaml)
    applies equally to the numbers shown here.
================================================================================
"""

import os
from datetime import datetime
from typing import Dict, List, Optional

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from engine.journals import JournalEntry

# ── Stallion brand styling (matches outputs/excel_exporter.py) ─────────────
STALLION_NAVY = RGBColor(0x1F, 0x38, 0x64)
STALLION_GOLD = RGBColor(0xC9, 0xA8, 0x4C)
GREY_TEXT     = RGBColor(0x59, 0x59, 0x59)
RED_TEXT      = RGBColor(0xC0, 0x00, 0x00)
WHITE_TEXT    = RGBColor(0xFF, 0xFF, 0xFF)

NAVY_HEX = "1F3864"
GREY_HEX = "F2F2F2"

GENERATED_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "outputs", "generated")


# ── Low-level formatting helpers ────────────────────────────────────────────

def _shade_cell(cell, hex_color: str) -> None:
    """Set a table cell's background fill — python-docx has no shading
    property directly, so this drops down to the underlying OOXML."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _money(value: Optional[float]) -> str:
    if value is None:
        return ""
    return f"GHS {value:,.2f}"


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    heading = doc.add_heading(level=level)
    run = heading.add_run(text)
    run.font.color.rgb = STALLION_NAVY
    run.font.bold = True
    run.font.size = Pt(18 if level == 1 else 14 if level == 2 else 12)


def _add_body(doc: Document, text: str, *, italic: bool = False, bold: bool = False, size: int = 10) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.italic = italic
    run.font.bold = bold


def _add_bullets(doc: Document, items: List[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        run.font.size = Pt(10)


def _add_kv_table(doc: Document, rows: List[tuple]) -> None:
    """Two-column label/value table — used for the assumptions and cover-page blocks."""
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.columns[0].width = Inches(2.2)
    table.columns[1].width = Inches(4.3)
    for label, value in rows:
        row = table.add_row().cells
        row[0].text = str(label)
        row[0].paragraphs[0].runs[0].font.bold = True
        row[0].paragraphs[0].runs[0].font.size = Pt(9)
        row[1].text = "" if value is None else str(value)
        if row[1].paragraphs[0].runs:
            row[1].paragraphs[0].runs[0].font.size = Pt(9)
    doc.add_paragraph()


def _add_data_table(
    doc: Document,
    headers: List[str],
    rows: List[list],
    *,
    total_row_indices: Optional[set] = None,
    onerous_row_indices: Optional[set] = None,
    font_size: int = 8,
) -> None:
    """Full data table with a navy header row (white bold text) — the Word
    equivalent of excel_exporter.py's _write_table()."""
    total_row_indices   = total_row_indices or set()
    onerous_row_indices = onerous_row_indices or set()

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    header_cells = table.rows[0].cells
    for col_idx, header in enumerate(headers):
        header_cells[col_idx].text = str(header)
        _shade_cell(header_cells[col_idx], NAVY_HEX)
        run = header_cells[col_idx].paragraphs[0].runs[0]
        run.font.color.rgb = WHITE_TEXT
        run.font.bold = True
        run.font.size = Pt(font_size)

    for r, row_data in enumerate(rows):
        cells = table.add_row().cells
        is_total   = r in total_row_indices
        is_onerous = r in onerous_row_indices
        for col_idx, value in enumerate(row_data):
            cells[col_idx].text = "" if value is None else str(value)
            if is_total:
                _shade_cell(cells[col_idx], GREY_HEX)
            if cells[col_idx].paragraphs[0].runs:
                run = cells[col_idx].paragraphs[0].runs[0]
                run.font.size = Pt(font_size)
                if is_total:
                    run.font.bold = True
                elif is_onerous and col_idx == 0:
                    run.font.color.rgb = RED_TEXT
                    run.font.bold = True

    doc.add_paragraph()


def _add_page_break(doc: Document) -> None:
    doc.add_page_break()


# ── Cover page ───────────────────────────────────────────────────────────────

def _build_cover_page(doc: Document, cover: dict) -> None:
    for _ in range(3):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(cover["title"])
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = STALLION_NAVY

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(cover["subtitle"])
    run.font.size = Pt(14)
    run.font.color.rgb = GREY_TEXT
    run.font.italic = True

    doc.add_paragraph()

    company = doc.add_paragraph()
    company.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = company.add_run(cover["company"])
    run.font.size = Pt(20)
    run.font.bold = True

    period = doc.add_paragraph()
    period.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = period.add_run(f"Reporting Period: {cover['period']}")
    run.font.size = Pt(12)

    for _ in range(4):
        doc.add_paragraph()

    details = [
        ("Report Date",         cover["report_date"]),
        ("Appointed Actuary",   cover["appointed_actuary"]),
        ("Consulting Firm",     cover["consulting_firm"]),
        ("Regulatory Basis",    cover["regulatory_basis"]),
        ("Reporting Standard",  cover["reporting_standard"]),
        ("Currency",            f"{cover['currency']}  ({cover['fx_rate']})"),
        ("Status",              cover["status"]),
    ]
    for label, value in details:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{label}:  ")
        run.font.bold = True
        run.font.size = Pt(10)
        run2 = p.add_run(str(value))
        run2.font.size = Pt(10)
        if label == "Status":
            run2.font.color.rgb = RED_TEXT
            run2.font.bold = True

    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Prepared using AMVS — Ghana Actuarial Modelling & Valuation System")
    run.font.size = Pt(8)
    run.font.color.rgb = GREY_TEXT
    run.font.italic = True

    _add_page_break(doc)


# ── Life-side sections (1-4, 6-8) — from avr dict ───────────────────────────

def _build_section1(doc: Document, s: dict) -> None:
    _add_heading(doc, s["title"])
    _add_body(doc, f"Reporting Period: {s['reporting_period']}  |  Frequency: {s['reporting_frequency']}", bold=True)
    _add_body(doc, f"Measurement Model: {s['measurement_model']}  |  In-force policies: {s['in_force_policies']:,}")
    doc.add_paragraph()
    _add_body(doc, "Key Findings:", bold=True)
    _add_bullets(doc, s["key_findings"])
    doc.add_paragraph()
    _add_body(doc, "Actuary's Comment:", bold=True)
    _add_body(doc, s["actuary_comment"], italic=True)


def _build_section2(doc: Document, s: dict) -> None:
    _add_heading(doc, s["title"])
    _add_kv_table(doc, [
        ("Product Name",      s["product_name"]),
        ("Product Type",      s["product_type"]),
        ("Coverage",          s["coverage"]),
        ("Benefit Tiers",     s["benefit_tiers"]),
        ("Premium Mode",      s["premium_mode"]),
        ("Policy Term",       s["policy_term"]),
        ("Entry Age Range",   s["entry_age_range"]),
        ("Portfolio Size",    f"{s['portfolio_size']:,}"),
        ("Cohort Year",       s["cohort_year"]),
        ("Portfolio Grouping",s["portfolio_grouping"]),
        ("Measurement Model", s["measurement_model"]),
    ])
    _add_body(doc, "Benefits Covered:", bold=True)
    _add_bullets(doc, s["benefits_covered"])


def _build_section3(doc: Document, s: dict) -> None:
    _add_heading(doc, s["title"])
    _add_body(doc, f"Basis Date: {s['basis_date']}", bold=True)

    _add_heading(doc, "Mortality", level=2)
    _add_kv_table(doc, list(s["mortality"].items()))

    _add_heading(doc, "Morbidity", level=2)
    _add_kv_table(doc, list(s["morbidity"].items()))

    _add_heading(doc, "Lapses", level=2)
    _add_kv_table(doc, list(s["lapses"].items()))

    _add_heading(doc, "Economic Assumptions", level=2)
    _add_kv_table(doc, list(s["economic"].items()))

    _add_heading(doc, "Expenses", level=2)
    _add_kv_table(doc, list(s["expenses"].items()))

    _add_heading(doc, "Commissions", level=2)
    _add_kv_table(doc, list(s["commissions"].items()))

    _add_heading(doc, "Premium Collection", level=2)
    _add_kv_table(doc, list(s["collection"].items()))

    _add_heading(doc, "Risk Adjustment", level=2)
    _add_kv_table(doc, list(s["risk_adjustment"].items()))

    _add_heading(doc, "Profit Target", level=2)
    _add_kv_table(doc, list(s["profit_target"].items()))


def _build_section4(doc: Document, s: dict) -> None:
    _add_heading(doc, s["title"])
    _add_kv_table(doc, [
        ("IFRS 17 Model",       s["ifrs17_model"]),
        ("PAA Eligibility",     s["paa_eligibility"]),
        ("Projection Method",   s["projection_method"]),
        ("Decrement Model",     s["decrement_model"]),
        ("Coverage Units",      s["coverage_units"]),
        ("CSM Release",         s["csm_release"]),
        ("Onerous Test",        s["onerous_test"]),
        ("Transition Approach", s["transition_approach"]),
        ("Reinsurance",         s["reinsurance"]),
    ])
    _add_body(doc, "Building Blocks:", bold=True)
    _add_bullets(doc, [s["building_blocks"]["block1"], s["building_blocks"]["block2"], s["building_blocks"]["block3"]])


def _build_section5(doc: Document, s: dict) -> None:
    _add_heading(doc, s["title"])
    _add_body(doc, f"Valuation Date: {s['valuation_date']}", bold=True)

    _add_heading(doc, "Liability for Remaining Coverage (LRC)", level=2)
    _add_data_table(doc, ["Component", "GHS"], [
        ["PVFCF",           _money(s["lrc"]["pvfcf"])],
        ["Risk Adjustment", _money(s["lrc"]["risk_adjustment"])],
        ["CSM",             _money(s["lrc"]["csm"])],
        ["Total LRC",       _money(s["lrc"]["total"])],
    ], total_row_indices={3})

    _add_heading(doc, "Liability for Incurred Claims (LIC)", level=2)
    _add_data_table(doc, ["Component", "GHS"], [
        ["Best Estimate",   _money(s["lic"]["best_estimate"])],
        ["Risk Adjustment", _money(s["lic"]["risk_adjustment"])],
        ["Total LIC",       _money(s["lic"]["total"])],
    ], total_row_indices={2})

    _add_data_table(doc, ["Metric", "GHS"], [
        ["Total Insurance Contract Liabilities", _money(s["total_liabilities"])],
        ["Total Liabilities (USD)",              f"USD {s['total_liabilities_usd']:,.0f}"],
    ], total_row_indices={0, 1})

    _add_heading(doc, "LRC Roll-Forward", level=2)
    lrc = s["lrc_rollforward"]
    _add_data_table(doc, ["Movement", "PVFCF", "RA", "CSM", "Total"], [
        ["Opening balance",     _money(lrc["opening_pvfcf"]), _money(lrc["opening_ra"]), _money(lrc["opening_csm"]), _money(lrc["opening_total"])],
        ["Premiums received",   _money(lrc["premiums_received"]), "", "", ""],
        ["Claims paid",         _money(lrc["claims_paid"]), "", "", ""],
        ["Finance income",      _money(lrc["finance_income"]), "", "", ""],
        ["CSM accretion",       "", "", _money(lrc["csm_accretion"]), ""],
        ["RA release",          "", _money(lrc["ra_release"]), "", ""],
        ["CSM amortisation",    "", "", _money(lrc["csm_amortisation"]), ""],
        ["Closing balance",     _money(lrc["closing_pvfcf"]), _money(lrc["closing_ra"]), _money(lrc["closing_csm"]), _money(lrc["closing_total"])],
    ], total_row_indices={0, 7})


def _build_section6(doc: Document, s: dict) -> None:
    _add_heading(doc, s["title"])
    _add_body(doc, f"Period: {s['period']}", bold=True)

    _add_heading(doc, "Insurance Revenue", level=2)
    r = s["insurance_revenue"]
    _add_data_table(doc, ["Component", "GHS"], [
        ["CSM Amortisation",  _money(r["csm_amortisation"])],
        ["RA Release",        _money(r["ra_release"])],
        ["Expected Claims",   _money(r["expected_claims"])],
        ["Experience Adj.",   _money(r["experience_adj"])],
        ["Total",             _money(r["total"])],
    ], total_row_indices={4})

    _add_heading(doc, "Insurance Expenses", level=2)
    e = s["insurance_expenses"]
    _add_data_table(doc, ["Component", "GHS"], [
        ["Incurred Claims",    _money(e["incurred_claims"])],
        ["Acquisition Costs",  _money(e["acquisition_costs"])],
        ["Other Expenses",     _money(e["other_expenses"])],
        ["Total",              _money(e["total"])],
    ], total_row_indices={3})

    _add_heading(doc, "Finance Income and Result", level=2)
    f = s["finance_income"]
    _add_data_table(doc, ["Metric", "GHS"], [
        ["Insurance Service Result",   _money(s["insurance_service_result"])],
        ["Finance Income (LRC)",       _money(f["lrc_interest"])],
        ["Finance Income (LIC)",       _money(f["lic_interest"])],
        ["Total Finance Income",       _money(f["total"])],
        ["Total Comprehensive Income", _money(s["total_comprehensive_income"])],
    ], total_row_indices={0, 4})


def _build_section7(doc: Document, s: dict) -> None:
    _add_heading(doc, s["title"])
    _add_body(doc, f"Period: {s['period']}  |  Locked-in rate: {s['locked_in_rate']}", bold=True)
    _add_data_table(doc, ["Movement", "GHS"], [
        ["Opening CSM",            _money(s["opening_csm"])],
        ["Interest Accretion",     _money(s["interest_accretion"])],
        ["Changes in Estimates",   _money(s["changes_estimates"])],
        ["CSM Amortisation",       _money(s["csm_amortisation"])],
        ["Closing CSM",            _money(s["closing_csm"])],
        ["Coverage Units (period)",f"{s['coverage_units_period']:,.0f}"],
        ["Coverage Units (total)", f"{s['coverage_units_total']:,.0f}"],
        ["Amortisation Rate",      f"{s['amortisation_rate']:.2f}%"],
    ], total_row_indices={0, 4})
    _add_body(doc, s["note"], italic=True)


def _build_section8(doc: Document, s: dict) -> None:
    _add_heading(doc, s["title"])
    _add_kv_table(doc, [
        ("Framework",           s["framework"]),
        ("Valuation Date",      s["valuation_date"]),
        ("Minimum CAR",         s["minimum_car"]),
        ("Solvency Status",     s["solvency_status"]),
    ])
    _add_data_table(doc, ["Metric", "GHS"], [
        ["Available Capital",   _money(s["available_capital"])],
        ["Required Capital",    _money(s["required_capital"])],
        ["Surplus / (Deficit)", _money(s["surplus_deficit"])],
        ["Capital Adequacy Ratio", f"{s['capital_adequacy_ratio']:.2f}%"],
    ], total_row_indices={0, 1})

    _add_heading(doc, "Risk Modules (Solvency Capital Requirement)", level=2)
    rm = s["risk_modules"]
    _add_data_table(doc, ["Risk Module", "GHS"], [
        ["Insurance Risk",   _money(rm["insurance_risk"])],
        ["Market Risk",      _money(rm["market_risk"])],
        ["Credit Risk",      _money(rm["credit_risk"])],
        ["Operational Risk", _money(rm["operational_risk"])],
        ["Total SCR",        _money(rm["total_scr"])],
    ], total_row_indices={4})

    _add_heading(doc, "Capital Composition", level=2)
    cc = s["capital_composition"]
    _add_data_table(doc, ["Capital Tier", "GHS"], [
        ["Tier 1 Capital",   _money(cc["tier1_capital"])],
        ["Tier 2 Capital",   _money(cc["tier2_capital"])],
        ["Total Available",  _money(cc["total_available"])],
    ], total_row_indices={2})


# ── Non-life section (5.5) — pulled directly from ifrs17_nonlife.py / journals.py ─

def _build_nonlife_section(
    doc: Document,
    statements: dict,
    journal_entries: List[JournalEntry],
    client_name: str,
) -> None:
    _add_heading(doc, "5.5 NON-LIFE CLAIMS RESERVES (GENERAL INSURANCE)")
    _add_body(
        doc,
        f"Client: {client_name}  |  Period: {statements['period']}  |  "
        f"RA loading: {statements['ra_loading']:.1%}  |  "
        f"Discounting: {'NIC RFR curve, ' + str(statements['discount_duration_years']) + ' yr assumed duration' if statements['discount_duration_years'] else 'Not applied'}",
        bold=True,
    )
    _add_body(
        doc,
        "PAA (Premium Allocation Approach) building blocks by class of business — "
        "LRC (Liability for Remaining Coverage = UPR - DAC) and LIC (Liability for "
        "Incurred Claims = IBNR + OCR + Effect of Discounting + ULAE + Risk Adjustment). "
        "See engine/ifrs17_nonlife.py for the documented modelling choices behind "
        "the discounting and risk adjustment figures.",
    )

    line_items = [
        ("IBNR + OCR",                              lambda c: c.ibnr + c.ocr),
        ("Effect of Discounting",                   lambda c: c.effect_of_discounting),
        ("Risk Adjustment",                         lambda c: c.risk_adjustment),
        ("ULAE",                                    lambda c: c.ulae),
        ("Liability for Incurred Claims (LIC)",     lambda c: c.lic),
        ("Unearned Premium Reserve",                lambda c: c.upr),
        ("Deferred Acquisition Cost",                lambda c: -c.dac),
        ("Liability for Remaining Coverage (LRC)",  lambda c: c.lrc),
        ("Total Reserve",                            lambda c: c.total_liability),
    ]
    subtotal_labels = {"Liability for Incurred Claims (LIC)", "Liability for Remaining Coverage (LRC)", "Total Reserve"}

    for basis in ("gross", "net", "ri"):
        _add_heading(doc, f"Balance Sheet — {basis.upper()}", level=2)
        headers = ["Line item"] + statements["classes"] + ["Total"]
        rows, total_idx = [], set()
        for i, (label, getter) in enumerate(line_items):
            values = [_money(getter(statements["by_class"][cls][basis])) for cls in statements["classes"]]
            total_val = _money(getter(statements["totals"][basis]))
            rows.append([label] + values + [total_val])
            if label in subtotal_labels:
                total_idx.add(i)
        _add_data_table(doc, headers, rows, total_row_indices=total_idx)

    _add_heading(doc, "Onerous Contract Test", level=2)
    onerous = [
        f"{cls} ({basis})" for cls in statements["classes"] for basis in ("gross", "net", "ri")
        if statements["by_class"][cls][basis].is_onerous
    ]
    _add_body(doc, "Onerous classes: " + (", ".join(onerous) if onerous else "None — no class is onerous"), bold=bool(onerous))

    _add_heading(doc, "Income Statement", level=2)

    def _pnl_net(cls: str, code: str, narrative_contains: Optional[str] = None) -> float:
        return round(sum((e.credit - e.debit) for e in journal_entries
                          if e.class_of_business == cls and e.account_code == code
                          and (narrative_contains is None or narrative_contains in e.narrative)), 2)

    # PIC's real chart of accounts combines claims, ULAE, risk adjustment
    # release, and acquisition cost recognition into ONE account, "204 - P&L
    # (PAA Insurance Expenses)" — see engine/journals.py's CHART_OF_ACCOUNTS.
    # Narrative is what tells these movement types apart here, not the code.
    headers = ["Line item"] + statements["classes"] + ["Total"]
    lines = [
        ("Insurance Revenue (premium earned)", "206", None),
        ("Claims Incurred",                     "204", "current service cost]"),
        ("ULAE",                                "204", "claims handling expenses"),
        ("Acquisition Costs (DAC deferral)",    "204", "Recognition of acquisition cost"),
    ]
    rows, service_result_row = [], {}
    for i, (label, code, narrative) in enumerate(lines):
        vals = [_pnl_net(cls, code, narrative) for cls in statements["classes"]]
        rows.append([label] + [_money(v) for v in vals] + [_money(round(sum(vals), 2))])
        service_result_row[i] = vals

    isr_vals = [sum(service_result_row[i][j] for i in range(len(lines))) for j in range(len(statements["classes"]))]
    rows.append(["Insurance Service Result"] + [_money(round(v, 2)) for v in isr_vals] + [_money(round(sum(isr_vals), 2))])
    isr_row_idx = len(rows) - 1

    ri_lines = [
        ("Finance Income (Effect of Discounting)", "205", None),
        ("Reinsurance Expense (Premium Ceded)",    "208", "Release of the reinsurance asset (Premium)"),
        ("Reinsurance Recoveries",                 "208", "Changes / Increase in recoverable amounts"),
    ]
    for label, code, narrative in ri_lines:
        vals = [_pnl_net(cls, code, narrative) for cls in statements["classes"]]
        rows.append([label] + [_money(v) for v in vals] + [_money(round(sum(vals), 2))])

    profit_vals = []
    for i, cls in enumerate(statements["classes"]):
        total = isr_vals[i]
        for label, code, narrative in ri_lines:
            total += _pnl_net(cls, code, narrative)
        profit_vals.append(round(total, 2))
    rows.append(["IFRS 17 Profit"] + [_money(v) for v in profit_vals] + [_money(round(sum(profit_vals), 2))])
    profit_row_idx = len(rows) - 1

    _add_data_table(doc, headers, rows, total_row_indices={isr_row_idx, profit_row_idx})

    _add_heading(doc, "Journal Entries", level=2)
    total_dr = round(sum(e.debit for e in journal_entries), 2)
    total_cr = round(sum(e.credit for e in journal_entries), 2)
    _add_body(
        doc,
        f"{len(journal_entries)} entries posted — Total Debit {_money(total_dr)} / "
        f"Total Credit {_money(total_cr)} ({'balanced' if abs(total_dr - total_cr) < 0.01 else 'NOT BALANCED'})",
        bold=True,
    )
    je_headers = ["Date", "Class", "Basis", "Acct", "Account Name", "Debit", "Credit", "Narrative"]
    je_rows = [
        [e.date, e.class_of_business, e.basis.upper(), e.account_code, e.account_name,
         _money(e.debit) if e.debit else "", _money(e.credit) if e.credit else "", e.narrative]
        for e in journal_entries
    ]
    je_rows.append(["", "", "", "", "TOTAL", _money(total_dr), _money(total_cr), ""])
    _add_data_table(doc, je_headers, je_rows, total_row_indices={len(je_rows) - 1}, font_size=7)


# ── Section 9 — Actuarial opinion and certificate (final page) ─────────────

def _build_certificate_page(doc: Document, s: dict) -> None:
    _add_page_break(doc)
    _add_heading(doc, s["title"])
    for paragraph in s["opinion"].split("\n\n"):
        _add_body(doc, paragraph.strip(), size=11)

    doc.add_paragraph()
    doc.add_paragraph()

    sig_line = doc.add_paragraph()
    sig_line.add_run("_" * 45).font.size = Pt(11)

    for label, value in [
        ("Signed:",         ""),
        ("Name:",           s["appointed_actuary"]),
        ("Qualification:",  s["qualifications"]),
        ("Firm:",           s["consulting_firm"]),
        ("Date:",           s["report_date"]),
    ]:
        p = doc.add_paragraph()
        run = p.add_run(f"{label}  ")
        run.font.bold = True
        run.font.size = Pt(11)
        run2 = p.add_run(str(value))
        run2.font.size = Pt(11)


# ── Top-level entry point ───────────────────────────────────────────────────

def generate_avr_word_document(
    avr:                       dict,
    nonlife_statements:          Optional[dict] = None,
    nonlife_journal_entries:       Optional[List[JournalEntry]] = None,
    nonlife_client_name:             Optional[str] = None,
    output_path:                       Optional[str] = None,
) -> str:
    """
    Build the complete NIC AVR Word document and write it to disk.

    Parameters:
        avr                       : api.nic_report.generate_avr_data() output
        nonlife_statements          : engine.ifrs17_nonlife.generate_nonlife_paa_statements()
                                       output — omit to skip the non-life section entirely
        nonlife_journal_entries       : engine.journals.generate_nonlife_journal() output
                                       (required if nonlife_statements is given)
        nonlife_client_name              : display name for the non-life section header
                                       (defaults to avr["cover"]["company"])
        output_path                        : file path to write to; defaults to
                                       outputs/generated/AVR_<company>_<period>_<timestamp>.docx

    Returns:
        The path the document was written to.
    """
    doc = Document()

    section = doc.sections[0]
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    _build_cover_page(doc, avr["cover"])
    _build_section1(doc, avr["section1"])
    _build_section2(doc, avr["section2"])
    _build_section3(doc, avr["section3"])
    _build_section4(doc, avr["section4"])
    _build_section5(doc, avr["section5"])

    if nonlife_statements is not None:
        _build_nonlife_section(
            doc, nonlife_statements, nonlife_journal_entries or [],
            nonlife_client_name or avr["cover"]["company"],
        )

    _build_section6(doc, avr["section6"])
    _build_section7(doc, avr["section7"])
    _build_section8(doc, avr["section8"])
    _build_certificate_page(doc, avr["section9"])

    if output_path is None:
        os.makedirs(GENERATED_DIR, exist_ok=True)
        company_slug = "".join(c if c.isalnum() else "_" for c in avr["cover"]["company"])[:40]
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_path = os.path.join(GENERATED_DIR, f"AVR_{company_slug}_{avr['cover']['period']}_{timestamp}.docx")

    doc.save(output_path)
    return output_path
