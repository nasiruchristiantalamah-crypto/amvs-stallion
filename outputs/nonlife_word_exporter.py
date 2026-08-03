"""
================================================================================
NON-LIFE NIC ACTUARIAL VALUATION REPORT — WORD (.docx) EXPORTER
================================================================================
What this file does:
    Builds the full non-life (PAA) Actuarial Valuation Report matching the
    exact section structure, table layouts, and column headers of PIC's
    own real published AVR ("2025 PIC Actuarial Valuation Report - Final
    (NIC Template).docx") — this is the report format Ghana's NIC actually
    expects for a general insurer, and it's what this project was built
    to reproduce (see README/session history). Reuses word_exporter.py's
    Stallion-branded styling helpers so both documents look like one system.

Data sources — every number below traces to a real engine computation:
    - engine.ifrs17_nonlife.generate_nonlife_paa_statements_granular() —
      LRC/LIC by the 6-class breakdown (Motor/Fire/Accident/Bonds/
      Engineering/Marine) PIC's real report uses.
    - engine.runner.run_nic_summary_granular() — the underlying IBNR/OCR/
      ULAE/UPR/DAC components, and load_paid_claims_granular_allocated()
      for cash claims paid.
    - engine.data_loader.load_triangle() — the raw claims triangles, for
      Appendix I.
    - data.yield_curve.load_yield_curve() — the discount curve, confirmed
      to match PIC's real published curve exactly.

Honest scope note — sections that CANNOT be computed from available data:
    - "Actual Versus Expected Analysis" (comparing this year's actual
      claims development against what was expected based on last year's
      valuation) needs a persisted PRIOR PERIOD closing position. This
      engine's non-life side computes a point-in-time snapshot, not a
      period-over-period roll-forward (see engine/journals.py's module
      docstring) — there's no prior-period snapshot to compare against.
      Flagged clearly in that section rather than fabricated.
    - Appendix IV/V (LRC/LIC and ARC/AIC reconciliation — opening balance
      -> movements -> closing balance) has the same limitation: without a
      persisted prior period, every movement is "day 1" first-time
      recognition (see engine/journals.py). Populated with what that
      implies, clearly labelled, not presented as a genuine multi-period
      roll-forward.
    - Company Overview (shareholders), Materiality Standards, and several
      other narrative sections describe facts about the business (share
      register, materiality threshold policy) that aren't actuarial
      calculations at all — no engine "computes" a shareholder list. These
      are templated with placeholders, parameterised by client name only;
      fill in real values via the `narrative` parameter for a real report.
================================================================================
"""

import os
from datetime import datetime
from typing import Dict, List, Optional

from docx import Document
from docx.shared import Inches, Pt

from data.yield_curve import load_yield_curve
from engine.clients import load_client
from engine.data_loader import RESERVING_CLASSES, load_triangle
from engine.ifrs17_nonlife import generate_nonlife_paa_statements_granular
from engine.journals import JournalEntry, generate_nonlife_journal
from engine.runner import load_paid_claims_granular_allocated, run_nic_summary_granular

from outputs.word_exporter import (
    GENERATED_DIR, RED_TEXT, STALLION_NAVY,
    _add_body, _add_bullets, _add_data_table, _add_heading, _add_kv_table, _add_page_break, _money,
)

REPORT_COLUMN_ORDER = ["Accident", "Bonds", "Engineering", "Fire", "Marine", "Motor"]
REPORT_COLUMN_DISPLAY_NAME = {"Fire": "Fire, Theft and Property", "Marine": "Marine and Aviation"}


def _dn(cls: str) -> str:
    return REPORT_COLUMN_DISPLAY_NAME.get(cls, cls)


def _thousands(value: float) -> str:
    """PIC's real report shows figures in GHS '000s throughout — matched here."""
    return f"{value / 1000:,.0f}"


# ── 1. Executive Summary ────────────────────────────────────────────────────

def _build_executive_summary(doc: Document, client_name: str, period: str, statements: dict, meta: dict) -> None:
    _add_heading(doc, "1. Executive Summary")

    _add_heading(doc, "1.1 Introduction", level=2)
    _add_body(doc,
        f"This report presents the actuarial valuation of {client_name}'s general insurance contract "
        f"liabilities as at {period}, prepared under IFRS 17 Insurance Contracts using the Premium "
        f"Allocation Approach (PAA), in accordance with the reporting requirements of the National "
        f"Insurance Commission (NIC) of Ghana.")

    _add_heading(doc, "1.2 Key Results Summary", level=2)
    t = statements["totals"]
    _add_data_table(doc, ["Reserve", "Gross (GHS '000)", "RI (GHS '000)", "Net (GHS '000)"], [
        ["Liability for Incurred Claims (LIC)", _thousands(t["gross"].lic), _thousands(t["ri"].lic), _thousands(t["net"].lic)],
        ["Liability for Remaining Coverage (LRC)", _thousands(t["gross"].lrc), _thousands(t["ri"].lrc), _thousands(t["net"].lrc)],
        ["Total Insurance Contract Liabilities", _thousands(t["gross"].total_liability), _thousands(t["ri"].total_liability), _thousands(t["net"].total_liability)],
    ], total_row_indices={2})

    _add_heading(doc, "1.3 Reserve Methodology", level=2)
    _add_data_table(doc, ["Insurance Liabilities", "Methodology"], [
        ["LIC — IBNR",           "Chain Ladder / Bornhuetter-Ferguson blend, credibility-weighted (see engine/bornhuetter_ferguson.py)"],
        ["LIC — Discounting",    "NIC published risk-free yield curve, single assumed average claim-payment duration"],
        ["LIC — Risk Adjustment","Percentage margin over best-estimate IBNR+OCR (simplified proxy — see note below)"],
        ["LRC — UPR / DAC",      "Client's own computed figures, read directly (not re-derived)"],
        ["LRC — Loss Component", "Onerous contract test: LRC < 0 triggers immediate loss recognition"],
    ])
    _add_body(doc,
        "Note: PIC's own Risk Adjustment methodology (Value at Risk, 75% confidence level, diversified "
        "portfolio basis) is a full statistical model this engine doesn't replicate — the figures above use "
        f"a simplified percentage-of-best-estimate proxy ({statements['ra_loading']:.0%}), disclosed here "
        "rather than presented as equivalent.", italic=True, size=9)

    _add_heading(doc, "1.4 Reserves Adequacy", level=2)
    onerous = [cls for cls in statements["classes"] if statements["by_class"][cls]["gross"].is_onerous]
    _add_body(doc, "No class of business is onerous — the Liability for Remaining Coverage is positive for every class."
              if not onerous else f"Onerous classes (negative LRC): {', '.join(onerous)}.", bold=bool(onerous))


# ── 2. Expression of Actuarial Opinion (narrative) ──────────────────────────

def _build_actuarial_opinion_section(doc: Document, client_name: str, period: str, actuary: str, firm: str) -> None:
    _add_heading(doc, "2. Expression of Actuarial Opinion")
    _add_body(doc,
        f"In my opinion, the general insurance contract liabilities of {client_name} as at {period}, "
        f"comprising the Liability for Incurred Claims and the Liability for Remaining Coverage, have "
        f"been determined in accordance with IFRS 17 Insurance Contracts and make reasonable and "
        f"adequate provision, in aggregate, for the company's outstanding obligations under its "
        f"insurance contracts as at the valuation date, subject to the assumptions, limitations, and "
        f"methodology disclosed throughout this report.")
    _add_body(doc, f"{actuary}, Appointed Actuary — {firm}", bold=True)


# ── 3-4. Company Overview / Materiality Standards (templated narrative) ────

def _build_company_overview(doc: Document, client_name: str, narrative: dict) -> None:
    _add_heading(doc, "3. Company's Overview")
    _add_body(doc, narrative.get("company_overview",
        f"{client_name} is a general (non-life) insurance company licensed and regulated by the National "
        f"Insurance Commission of Ghana. [Company overview narrative — ownership structure, principal "
        f"activities, licensed classes of business — to be supplied per client; not an actuarial "
        f"calculation this engine produces.]"))
    if narrative.get("shareholders"):
        _add_data_table(doc, ["Shareholder", "% Shareholding"], [
            [s["name"], s["percentage"]] for s in narrative["shareholders"]
        ])


def _build_materiality_standards(doc: Document, narrative: dict) -> None:
    _add_heading(doc, "4. Materiality Standards")
    _add_body(doc, narrative.get("materiality",
        "[Materiality threshold policy — a governance/judgement decision the Appointed Actuary and "
        "company set together, not an actuarial calculation this engine derives. Supply per client via "
        "the `narrative` parameter.]"))


# ── 5. Data ──────────────────────────────────────────────────────────────────

def _build_data_section(doc: Document, statements: dict, journal_entries: List[JournalEntry], period: str) -> None:
    _add_heading(doc, "5. Data")
    _add_body(doc,
        "This valuation is based on data extracted directly from the company's own claims, premium, and "
        "policy administration systems for the period, read and reconciled programmatically (see "
        "engine/data_loader.py) rather than manually re-keyed.")

    classes = [c for c in REPORT_COLUMN_ORDER if c in statements["classes"]] or statements["classes"]

    _add_heading(doc, "5.1 Gross Written Premium", level=3)
    upr_by_class = statements["reserving_summary"]["by_class"]
    _add_data_table(doc, ["Class of Business"] + [_dn(c) for c in classes] + ["Total"],
        [["GWP (GHS '000)"] + [_thousands(upr_by_class[c]["gross"]["upr"]) for c in classes] +
         [_thousands(sum(upr_by_class[c]["gross"]["upr"] for c in classes))]],
    )
    _add_body(doc, "Note: shown here as Unearned Premium Reserve (the client's own computed figure); "
              "Written Premium by class is available from the same source workbook on request.", italic=True, size=9)

    _add_heading(doc, "5.2 Gross Claims Paid", level=3)
    paid = load_paid_claims_granular_allocated(_client_id_from_statements(statements))
    _add_data_table(doc, ["Class of Business"] + [_dn(c) for c in classes] + ["Total"],
        [["Paid Claims (GHS '000)"] + [_thousands(paid.get(c, 0.0)) for c in classes] +
         [_thousands(sum(paid.get(c, 0.0) for c in classes))]],
    )

    _add_heading(doc, "5.3 Gross Outstanding Claims (OCR)", level=3)
    _add_data_table(doc, ["Class of Business"] + [_dn(c) for c in classes] + ["Total"],
        [["OCR (GHS '000)"] + [_thousands(upr_by_class[c]["gross"]["ocr"]) for c in classes] +
         [_thousands(sum(upr_by_class[c]["gross"]["ocr"] for c in classes))]],
    )

    _add_heading(doc, "5.4 Data Review, Sufficiency, and Reliability", level=3)
    _add_body(doc,
        "Data was read directly from the client's own source workbooks using structural parsing (header "
        "text matching rather than fixed cell references), validated against the client's own published "
        "figures class-by-class before being used in this valuation — see engine/data_loader.py and the "
        "project's test suite for the validation methodology. No manual adjustments were made to the "
        "underlying data.")

    _add_heading(doc, "5.5 Reliance and Limitations", level=3)
    _add_body(doc,
        "This valuation relies on the completeness and accuracy of data supplied by the company. No "
        "independent audit of the underlying policy or claims administration systems was performed. "
        "Known modelling simplifications (Risk Adjustment methodology, discounting duration assumption, "
        "and the absence of a persisted prior-period roll-forward for the Actual vs Expected and "
        "reconciliation sections) are disclosed throughout this report rather than presented as exact.",
        italic=True)


def _client_id_from_statements(statements: dict) -> str:
    # statements doesn't carry client_id directly — callers pass it through
    # the module-level generate function instead; this is a narrow internal
    # helper used only within this file's own section builders.
    return statements.get("_client_id", "pic")


# ── 6-7. Expenses / Portfolio Reporting (templated narrative) ──────────────

def _build_expenses_section(doc: Document, statements: dict) -> None:
    _add_heading(doc, "6. Expenses")
    ulae_by_class = statements["reserving_summary"]["by_class"]
    classes = [c for c in REPORT_COLUMN_ORDER if c in statements["classes"]] or statements["classes"]
    _add_data_table(doc, ["Class of Business"] + [_dn(c) for c in classes] + ["Total"],
        [["ULAE (GHS '000)"] + [_thousands(ulae_by_class[c]["gross"]["ulae"]) for c in classes] +
         [_thousands(sum(ulae_by_class[c]["gross"]["ulae"] for c in classes))]],
    )
    _add_body(doc, "Unallocated Loss Adjustment Expense (ULAE), computed by the client using the "
              "Alpha-ratio method and read directly — not reinsured, so Net = Gross for this line.", italic=True, size=9)


def _build_portfolio_reporting(doc: Document, narrative: dict) -> None:
    _add_heading(doc, "7. Portfolio Reporting")
    _add_heading(doc, "7.1 Determination of Portfolios", level=2)
    _add_body(doc, narrative.get("portfolios",
        "Insurance contracts are grouped into portfolios by class of business, consistent with how the "
        "company prices, manages, and reports on its risks internally — see Appendix II for the mapping "
        "of underlying products to IFRS 17 portfolios."))
    _add_heading(doc, "7.2 Reinsurance Contracts Held", level=2)
    _add_body(doc, narrative.get("reinsurance",
        "[Reinsurance treaty structure — quota share / surplus / excess-of-loss terms per class. See "
        "clients/<id>/assumptions.yaml's reinsurance: section for what's configured; ceded figures "
        "throughout this report are read directly from the client's own RI workbooks, not modelled from "
        "treaty terms.]"))
    _add_heading(doc, "7.3 Profitability Groupings", level=2)
    _add_body(doc, "Contracts are assessed for onerousness at initial recognition and at each reporting "
              "date at the class-of-business level (see Section 1.4) — no contract group within scope was "
              "identified as onerous at the valuation date unless explicitly flagged.")


# ── 8. Discount Curve ────────────────────────────────────────────────────────

def _build_discount_curve_section(doc: Document) -> None:
    _add_heading(doc, "8. Discount Curve")
    _add_body(doc, "The National Insurance Commission's published risk-free yield curve, used to discount "
              "the estimates of future cash flows for both LIC and LRC throughout this report.")
    curve = load_yield_curve()
    years = sorted(curve.keys())[:15]
    _add_data_table(doc, ["Year", "Discount Rate (%)"],
        [[y, f"{curve[y]:.2%}"] for y in years], font_size=9)


# ── 9. Risk Adjustment (narrative) ──────────────────────────────────────────

def _build_risk_adjustment_section(doc: Document, statements: dict) -> None:
    _add_heading(doc, "9. Risk Adjustment for Non-Financial Risk")
    _add_body(doc,
        f"Risk Adjustment is computed as a {statements['ra_loading']:.0%} loading over the best-estimate "
        f"(undiscounted IBNR + OCR) — a simplified percentage-margin proxy. PIC's own real methodology "
        f"(Value at Risk, 75% confidence level, estimated at a diversified/total-portfolio level) is a "
        f"full statistical model this engine doesn't replicate; disclosed here rather than presented as "
        f"equivalent.", italic=True)


# ── 10. Liability for Incurred Claims ───────────────────────────────────────

def _build_lic_section(doc: Document, statements: dict) -> None:
    _add_heading(doc, "10. Liability for Incurred Claims")
    _add_heading(doc, "10.1 Estimates of Future Cash Flows for LIC", level=2)
    classes = [c for c in REPORT_COLUMN_ORDER if c in statements["classes"]] or statements["classes"]

    for basis, title in (("gross", "Liability for Incurred Claims (Gross)"), ("ri", "Assets for Incurred Claims (Reinsurance)")):
        headers = ["Component"] + [_dn(c) for c in classes] + ["Total"]
        rows = []
        for label, getter in [
            ("Best Estimate (IBNR+OCR+ULAE)", lambda c: c.ibnr + c.ocr + c.ulae),
            ("Risk Adjustment",                lambda c: c.risk_adjustment),
            ("Total LIC",                       lambda c: c.lic),
        ]:
            vals = [getter(statements["by_class"][cls][basis]) for cls in classes]
            rows.append([label] + [_thousands(v) for v in vals] + [_thousands(sum(vals))])
        _add_heading(doc, title, level=3)
        _add_data_table(doc, headers, rows, total_row_indices={2})

    _add_heading(doc, "10.2 Discounting the Estimates of Future Cash Flows", level=2)
    headers = ["Class of Business"] + [_dn(c) for c in classes] + ["Total"]
    vals = [statements["by_class"][cls]["gross"].effect_of_discounting for cls in classes]
    _add_data_table(doc, headers, [["Effect of Discounting (GHS '000)"] + [_thousands(v) for v in vals] + [_thousands(sum(vals))]])


# ── 11. Actual vs Expected Analysis — not available ─────────────────────────

def _build_ave_section(doc: Document) -> None:
    _add_heading(doc, "11. Actual Versus Expected Analysis for Prior Year-End Valuations")
    _add_body(doc,
        "Not available. This analysis compares this period's actual claims development against what "
        "was expected based on the prior period's valuation — it requires a PERSISTED prior-period "
        "closing position. This engine's non-life side currently computes a point-in-time snapshot each "
        "run, not a period-over-period roll-forward (unlike the life side — see "
        "engine/rollforward_store.py), so there is no prior-period baseline available to compare "
        "against. This is a disclosed gap, not a fabricated result.", bold=True)


# ── 12. Liability for Remaining Coverage ────────────────────────────────────

def _build_lrc_section(doc: Document, statements: dict) -> None:
    _add_heading(doc, "12. Liability for Remaining Coverage (LRC)")
    _add_heading(doc, "12.1 Measurement Approach", level=2)
    _add_body(doc, "Premium Allocation Approach (PAA) — LRC = Unearned Premium Reserve less Deferred "
              "Acquisition Costs, with a loss component recognised where the balance would be negative "
              "(onerous contract test).")

    classes = [c for c in REPORT_COLUMN_ORDER if c in statements["classes"]] or statements["classes"]
    _add_heading(doc, "12.2 Estimates of Future Cash Flows", level=2)
    for basis, title in (("gross", "Liability for Remaining Coverage (Gross)"), ("ri", "Assets for Remaining Coverage (Reinsurance)")):
        headers = ["Component"] + [_dn(c) for c in classes] + ["Total"]
        rows = []
        for label, getter in [
            ("Unearned Premium Reserve",   lambda c: c.upr),
            ("Deferred Acquisition Cost",  lambda c: -c.dac),
            ("Loss Component",              lambda c: c.loss_component),
            ("Total LRC",                    lambda c: c.lrc),
        ]:
            vals = [getter(statements["by_class"][cls][basis]) for cls in classes]
            rows.append([label] + [_thousands(v) for v in vals] + [_thousands(sum(vals))])
        _add_heading(doc, title, level=3)
        _add_data_table(doc, headers, rows, total_row_indices={3})


# ── 13. Summary of Results ──────────────────────────────────────────────────

def _build_summary_of_results(doc: Document, statements: dict, prior_period_note: bool = True) -> None:
    _add_heading(doc, "13. Summary of Results")
    t = statements["totals"]
    _add_data_table(doc, ["Reserve", "Current Period (GHS '000)", "Basis"], [
        ["Liability for Incurred Claims (LIC)",            _thousands(t["gross"].lic), "Gross"],
        ["Liability for Remaining Coverage (LRC)",          _thousands(t["gross"].lrc), "Gross"],
        ["Total Gross Insurance Liabilities",                 _thousands(t["gross"].total_liability), "Gross"],
        ["Reinsurance Assets",                                  _thousands(t["ri"].total_liability), "RI"],
        ["Net Insurance Liabilities",                             _thousands(t["net"].total_liability), "Net"],
    ], total_row_indices={2, 4})
    if prior_period_note:
        _add_body(doc, "Prior-period comparative figures are not shown — see Section 11 for why a "
                  "period-over-period comparison isn't currently available from this engine.", italic=True, size=9)


# ── 14. Definitions ──────────────────────────────────────────────────────────

DEFINITIONS = [
    ("IBNR",  "Incurred But Not Reported — the estimated liability for claims that have occurred but not yet been reported to the insurer."),
    ("OCR",   "Outstanding Claims Reserve — case-by-case reserves for reported, not-yet-settled claims."),
    ("ULAE",  "Unallocated Loss Adjustment Expense — the cost of administering and settling claims, not attributable to any single claim."),
    ("UPR",   "Unearned Premium Reserve — the portion of written premium relating to coverage not yet provided."),
    ("DAC",   "Deferred Acquisition Cost — acquisition costs capitalised and released over the coverage period, matching the UPR release pattern."),
    ("LIC",   "Liability for Incurred Claims — IBNR + OCR + ULAE + Risk Adjustment, discounted."),
    ("LRC",   "Liability for Remaining Coverage — UPR less DAC, plus any Loss Component."),
    ("PAA",   "Premium Allocation Approach — the simplified IFRS 17 measurement model available for contracts of one year or less (or where it doesn't materially differ from the General Measurement Model)."),
    ("Risk Adjustment", "The compensation an insurer requires for bearing the uncertainty in the amount and timing of cash flows from non-financial risk."),
    ("URR",   "Unexpired Risk Reserve — the portion of an immature underwriting year's claims development attributable to risk that hasn't occurred yet (the policy term hasn't run its course) — conceptually distinct from, and excluded from, IBNR to avoid double-counting risk already provided for via UPR."),
]


def _build_definitions_section(doc: Document) -> None:
    _add_heading(doc, "14. Definitions")
    _add_data_table(doc, ["Term", "Definition"], [[t, d] for t, d in DEFINITIONS], font_size=9)


# ── Appendix I: Claims Data Triangles ───────────────────────────────────────

def _build_appendix_i(doc: Document, client_id: str, data_folder_override: Optional[str] = None) -> None:
    _add_page_break(doc)
    _add_heading(doc, "Appendix I: Claims Data Triangles")
    for cls in RESERVING_CLASSES:
        tri = load_triangle(cls, client_id=client_id, data_folder_override=data_folder_override)
        gross_tri = tri["gross_triangle"]
        _add_heading(doc, f"{cls} — Gross Cumulative Incurred Claims", level=2)
        origin_years = sorted(gross_tri.keys())
        max_dev = max(len(v) for v in gross_tri.values())
        headers = ["Underwriting Year"] + [f"Dev. Yr {i}" for i in range(max_dev)]
        rows = [[oy] + [_thousands(v) if v else "" for v in gross_tri[oy]] + [""] * (max_dev - len(gross_tri[oy]))
                for oy in origin_years]
        _add_data_table(doc, headers, rows, font_size=7)


# ── Appendix II: Aggregation of portfolios (templated) ──────────────────────

def _build_appendix_ii(doc: Document, narrative: dict) -> None:
    _add_heading(doc, "Appendix II: Aggregation of Insurance Contracts into IFRS 17 Portfolios")
    _add_body(doc, narrative.get("portfolio_mapping",
        "[Mapping of underlying products/classes to IFRS 17 portfolios — a governance/product-structure "
        "decision, not an actuarial calculation. Supply per client via the `narrative` parameter.]"))


# ── Appendix III: Breakdown of Undiscounted Liabilities ─────────────────────

def _build_appendix_iii(doc: Document, statements: dict) -> None:
    _add_page_break(doc)
    _add_heading(doc, "Appendix III: Breakdown of Undiscounted Liabilities")
    classes = [c for c in REPORT_COLUMN_ORDER if c in statements["classes"]] or statements["classes"]
    for basis, title in (("gross", "Gross"), ("ri", "Reinsurance (ceded)")):
        headers = ["Class of business", "IBNR", "OCR", "ULAE", "UPR", "DAC"]
        rows, total_idx = [], set()
        totals = {m: 0.0 for m in ("ibnr", "ocr", "ulae", "upr", "dac")}
        for cls in classes:
            c = statements["by_class"][cls][basis]
            vals = [c.ibnr, c.ocr, c.ulae, c.upr, c.dac]
            for m, v in zip(totals, vals):
                totals[m] += v
            rows.append([cls] + [_thousands(v) for v in vals])
        rows.append(["Total"] + [_thousands(v) for v in totals.values()])
        total_idx.add(len(rows) - 1)
        _add_heading(doc, title, level=2)
        _add_data_table(doc, headers, rows, total_row_indices=total_idx)


# ── Appendix IV/V: Reconciliations (day-1 basis, disclosed) ────────────────

def _build_reconciliation_appendix(doc: Document, title: str, statements: dict, journal_entries: List[JournalEntry], basis: str) -> None:
    _add_page_break(doc)
    _add_heading(doc, title)
    _add_body(doc,
        "Shown on a 'day 1' (first-time recognition) basis — the opening balance is nil and the full "
        "current balance appears as this period's movement, since this engine doesn't yet persist a "
        "prior-period non-life closing position to roll forward from (see Section 11). A genuine "
        "multi-period reconciliation would show a non-zero opening balance carried from the prior "
        "valuation.", italic=True)

    classes = [c for c in REPORT_COLUMN_ORDER if c in statements["classes"]] or statements["classes"]
    headers = ["Class of business", "LRC (excl. Loss Component)", "Loss Component", "LIC (incl. RA)", "Total"]
    rows, total_idx = [], set()
    tot = [0.0, 0.0, 0.0, 0.0]
    for cls in classes:
        c = statements["by_class"][cls][basis]
        lrc_excl_lc = c.lrc + c.loss_component if c.is_onerous else c.lrc
        vals = [lrc_excl_lc, c.loss_component, c.lic, lrc_excl_lc - c.loss_component + c.lic]
        for i, v in enumerate(vals):
            tot[i] += v
        rows.append([cls] + [_thousands(v) for v in vals])
    rows.append(["Total"] + [_thousands(v) for v in tot])
    total_idx.add(len(rows) - 1)
    _add_data_table(doc, headers, rows, total_row_indices=total_idx)


# ── Appendix VI/VII: Actuarial Opinion / Certificate of Data Accuracy ──────

def _build_certificate_pages(doc: Document, client_name: str, period: str, actuary: str, firm: str, qualifications: str, report_date: str) -> None:
    _add_page_break(doc)
    _add_heading(doc, "Appendix VI: Actuarial Opinion")
    _add_body(doc,
        f"I, {actuary}, being the Appointed Actuary of {client_name}, hereby certify that the general "
        f"insurance contract liabilities of {client_name} as at {period} have been determined in "
        f"accordance with IFRS 17 Insurance Contracts and the requirements of the National Insurance "
        f"Commission of Ghana, using the Premium Allocation Approach, and that the assumptions and "
        f"methodology applied are appropriate for this purpose, subject to the limitations disclosed "
        f"throughout this report.", size=11)
    doc.add_paragraph()
    doc.add_paragraph()
    sig = doc.add_paragraph()
    sig.add_run("_" * 45).font.size = Pt(11)
    for label, value in [("Signed:", ""), ("Name:", actuary), ("Qualification:", qualifications),
                          ("Firm:", firm), ("Date:", report_date)]:
        p = doc.add_paragraph()
        p.add_run(f"{label}  ").font.bold = True
        p.add_run(str(value)).font.size = Pt(11)

    _add_page_break(doc)
    _add_heading(doc, "Appendix VII: Certificate of Data Accuracy")
    _add_body(doc,
        f"{client_name} confirms that the data supplied for this valuation is, to the best of its "
        f"knowledge, complete and accurate, and that any known data limitations have been disclosed to "
        f"the Appointed Actuary.", size=11)
    doc.add_paragraph()
    doc.add_paragraph()
    sig2 = doc.add_paragraph()
    sig2.add_run("_" * 45).font.size = Pt(11)
    p2 = doc.add_paragraph()
    p2.add_run("Signed (on behalf of ").font.size = Pt(11)
    p2.add_run(client_name).font.bold = True
    p2.add_run("):").font.size = Pt(11)


# ── Cover page ───────────────────────────────────────────────────────────────

def _build_cover_page(doc: Document, client_name: str, period: str, report_date: str, actuary: str, firm: str) -> None:
    for _ in range(3):
        doc.add_paragraph()
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("ACTUARIAL VALUATION REPORT")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = STALLION_NAVY

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("General Insurance (Non-Life) — IFRS 17 Premium Allocation Approach")
    run.font.size = Pt(14)
    run.font.italic = True

    doc.add_paragraph()
    company = doc.add_paragraph()
    company.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = company.add_run(client_name)
    run.font.size = Pt(20)
    run.font.bold = True

    for _ in range(4):
        doc.add_paragraph()
    for label, value in [("Reporting Period", period), ("Report Date", report_date),
                          ("Appointed Actuary", actuary), ("Consulting Firm", firm),
                          ("Regulatory Basis", "National Insurance Commission (NIC) — Ghana"),
                          ("Status", "DRAFT FOR REVIEW")]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(f"{label}:  ").font.bold = True
        p.add_run(str(value))
    _add_page_break(doc)


# ── Top-level entry point ───────────────────────────────────────────────────

def generate_nonlife_avr_word_document(
    client_id:            str            = "pic",
    period:                str            = "FY2025",
    appointed_actuary:      str            = "Charles Osei-Akoto",
    consulting_firm:          str            = "Stallion Consultants Ltd",
    qualifications:             str            = "Fellow, Institute and Faculty of Actuaries (FIA)",
    narrative:                     Optional[dict] = None,
    output_path:                       Optional[str] = None,
    data_folder_override:               Optional[str] = None,
    company_name_override:                Optional[str] = None,
) -> str:
    """
    Build the complete non-life NIC AVR Word document, matching PIC's real
    published report structure, and write it to disk.

    Parameters:
        client_id           : which client (clients/<client_id>/)
        period               : reporting period label, e.g. "FY2025"
        narrative              : optional dict of client-specific narrative
                                 text overrides for the non-computed
                                 sections (company_overview, shareholders,
                                 materiality, portfolios, reinsurance,
                                 portfolio_mapping) — see each section
                                 builder's docstring for keys. Sensible
                                 placeholder text is used for anything omitted.
        output_path                : file path to write to; defaults to
                                 outputs/generated/AVR_NonLife_<client>_<period>_<timestamp>.docx

    Returns:
        The path the document was written to.
    """
    narrative = narrative or {}
    client = load_client(client_id)
    display_name = company_name_override or client.name
    report_date = datetime.now().strftime("%d %B %Y")

    statements = generate_nonlife_paa_statements_granular(
        client_id=client_id, period=period, data_folder_override=data_folder_override, verbose=False,
    )
    statements["_client_id"] = client_id
    paid = load_paid_claims_granular_allocated(client_id=client_id, data_folder_override=data_folder_override)
    journal_entries = generate_nonlife_journal(statements, paid, period=period)

    doc = Document()
    section = doc.sections[0]
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    _build_cover_page(doc, display_name, period, report_date, appointed_actuary, consulting_firm)
    _build_executive_summary(doc, display_name, period, statements, {})
    _build_actuarial_opinion_section(doc, display_name, period, appointed_actuary, consulting_firm)
    _build_company_overview(doc, display_name, narrative)
    _build_materiality_standards(doc, narrative)
    _build_data_section(doc, statements, journal_entries, period)
    _build_expenses_section(doc, statements)
    _build_portfolio_reporting(doc, narrative)
    _build_discount_curve_section(doc)
    _build_risk_adjustment_section(doc, statements)
    _build_lic_section(doc, statements)
    _build_ave_section(doc)
    _build_lrc_section(doc, statements)
    _build_summary_of_results(doc, statements)
    _build_definitions_section(doc)
    _build_appendix_i(doc, client_id, data_folder_override)
    _build_appendix_ii(doc, narrative)
    _build_appendix_iii(doc, statements)
    _build_reconciliation_appendix(doc, "Appendix IV: Reconciliation of the Liability for Remaining Coverage & Incurred Claims", statements, journal_entries, "gross")
    _build_reconciliation_appendix(doc, "Appendix V: Reconciliation of the Asset for Remaining Coverage and Incurred Claims", statements, journal_entries, "ri")
    _build_certificate_pages(doc, display_name, period, appointed_actuary, consulting_firm, qualifications, report_date)

    if output_path is None:
        os.makedirs(GENERATED_DIR, exist_ok=True)
        client_slug = "".join(c if c.isalnum() else "_" for c in display_name)[:40]
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_path = os.path.join(GENERATED_DIR, f"AVR_NonLife_{client_slug}_{period}_{timestamp}.docx")

    doc.save(output_path)
    return output_path
