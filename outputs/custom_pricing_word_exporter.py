"""
================================================================================
CUSTOM PRODUCT PRICING — AVR product note (Word)
================================================================================
What this file does:
    Writes the "Generate AVR Note" button's Word document for the
    dashboard's Part 4 product pricing platform — a short actuarial product
    note documenting the product definition, the pricing basis used, and
    the resulting premium/IFRS 17 figures (engine/custom_pricing.py's
    run_custom_pricing() output). This is a lighter-weight, single-product
    note — not the full multi-section NIC AVR report structure
    outputs/nonlife_word_exporter.py builds for the non-life reserving side.
================================================================================
"""

import os
from datetime import datetime
from typing import List, Optional

from docx import Document
from docx.shared import Inches, Pt

GENERATED_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "outputs", "generated")


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def _add_body(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(8)


def _add_kv_table(doc: Document, rows: List[tuple]) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.style = "Light Grid Accent 1"
    for k, v in rows:
        row = table.add_row().cells
        row[0].text = str(k)
        row[1].text = str(v)


def generate_custom_pricing_avr_note(
    result:            dict,
    product_spec:      Optional[dict] = None,
    consulting_firm:   str = "Stallion Consultants Ltd",
    appointed_actuary: str = "",
    output_path:       Optional[str] = None,
) -> str:
    """
    Build a short actuarial product note documenting one custom-priced
    product's basis and results, and write it to disk.

    Returns:
        The path the document was written to.
    """
    product_spec = product_spec or {}
    report_date = datetime.now().strftime("%d %B %Y")
    a = result.get("assumptions_used", {})

    doc = Document()
    section = doc.sections[0]
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    doc.add_heading(f"Actuarial Product Note — {result.get('product_name', 'Custom Product')}", level=0)
    _add_body(doc, f"Prepared by {consulting_firm}" + (f" — {appointed_actuary}" if appointed_actuary else ""))
    _add_body(doc, f"Date: {report_date}")

    _add_heading(doc, "1. Product Description")
    _add_kv_table(doc, [
        ("Product type", result.get("product_type")),
        ("Policy term", product_spec.get("policy_term_years") or "Whole of life"),
        ("Premium mode", product_spec.get("premium_mode", "monthly")),
        ("Sum assured (GHS)", product_spec.get("sum_assured")),
        ("Entry age", product_spec.get("entry_age")),
        ("Gender basis", product_spec.get("gender", "unisex")),
    ])

    riders = product_spec.get("riders") or []
    if riders:
        _add_heading(doc, "2. Benefit Riders", level=2)
        table = doc.add_table(rows=1, cols=3)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text = "Rider", "Benefit type", "Benefit amount (GHS)"
        for r in riders:
            row = table.add_row().cells
            row[0].text = str(r.get("name", ""))
            row[1].text = str(r.get("benefit_type", ""))
            row[2].text = str(r.get("benefit_amount", ""))

    _add_heading(doc, "3. Pricing Basis")
    _add_body(doc, f"Basis used: {a.get('name', 'Ghana Market Defaults')}")
    _add_kv_table(doc, [
        ("Mortality loading", a.get("mortality_loading")),
        ("Gender basis", a.get("gender_main_str")),
        ("Valuation rate (p.a.)", a.get("valuation_rate_pa")),
        ("Investment return (p.a.)", a.get("investment_rate_pa")),
        ("Collection rate", a.get("collection_rate")),
        ("Target profit margin", a.get("target_profit_margin")),
        ("Acquisition cost (GHS)", a.get("acquisition_cost")),
        ("Renewal expense (GHS/year)", a.get("renewal_expense_annual")),
    ])

    _add_heading(doc, "4. Pricing Results")
    _add_kv_table(doc, [
        ("Monthly premium (GHS)", result.get("monthly_premium")),
        ("Annual premium (GHS)", result.get("annual_premium")),
        ("Single premium equivalent (GHS)", result.get("single_premium_equiv")),
        ("Profit margin achieved", result.get("profit_margin")),
    ])

    _add_heading(doc, "5. IFRS 17 at Inception")
    _add_kv_table(doc, [
        ("PVFCF (GHS)", result.get("pvfcf")),
        ("Risk Adjustment (GHS)", result.get("risk_adjustment")),
        ("CSM at inception (GHS)", result.get("csm_at_inception")),
        ("LRC total (GHS)", result.get("lrc_total")),
        ("Onerous contract?", "Yes" if result.get("is_onerous") else "No"),
        ("Loss component (GHS)", result.get("loss_component")),
    ])

    breakeven = (result.get("profit_signature") or {}).get("breakeven_year")
    _add_heading(doc, "6. Profit Emergence")
    _add_body(doc, f"Breakeven is projected in policy year {breakeven}." if breakeven else "Breakeven is not reached within the projection period.")

    if output_path is None:
        os.makedirs(GENERATED_DIR, exist_ok=True)
        name_slug = "".join(c if c.isalnum() else "_" for c in result.get("product_name", "CustomProduct"))[:40]
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_path = os.path.join(GENERATED_DIR, f"AVRNote_{name_slug}_{timestamp}.docx")

    doc.save(output_path)
    return output_path
