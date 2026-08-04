"""
================================================================================
GIRBC SOLVENCY CERTIFICATE — WORD (.docx) EXPORTER
================================================================================
What this file does:
    Builds a formatted GIRBC solvency certificate matching NIC requirements:
    company name, valuation date, GIRBC CAR, legacy CAR, solvency status,
    risk-module breakdown, capital composition, 5 prescribed stress test
    results, an actuarial opinion paragraph, and a signature block —
    reusing outputs/word_exporter.py's shared Stallion-branded styling
    helpers, exactly as outputs/nonlife_word_exporter.py does for the
    non-life AVR (same pattern, separate file, not crammed into
    word_exporter.py itself since that file is specifically the life-side
    AVR exporter).

    Self-contained like the other exporters here: given a client_id, this
    loads that client's real GIRBC/legacy workbooks itself
    (engine.data_loader.load_rbc_solvency_data) and runs the full
    engine/rbc/ pipeline internally — no pre-computed results need to be
    passed in, though generate_girbc_certificate_from_results() is also
    exposed for the API layer to reuse results it already computed (e.g.
    from a user's manually-edited exposure inputs) instead of paying for a
    second calculation.
================================================================================
"""

import os
from datetime import datetime
from typing import List, Optional

from docx import Document
from docx.shared import Pt

from engine.clients import load_client
from engine.data_loader import load_rbc_solvency_data
from engine.rbc.aggregation import SolvencyResult, calculate_solvency
from engine.rbc.credit_risk import calculate_credit_risk
from engine.rbc.data_model import QualifyingCapitalResources
from engine.rbc.insurance_risk import calculate_insurance_risk
from engine.rbc.legacy_solvency import LegacySolvencyResult, calculate_legacy_solvency
from engine.rbc.market_risk import calculate_market_risk
from engine.rbc.operational_risk import calculate_operational_risk
from engine.rbc.stress_tests import StressTestResult, run_stress_tests

from outputs.word_exporter import (
    GENERATED_DIR, RED_TEXT, STALLION_NAVY,
    _add_body, _add_data_table, _add_heading, _add_kv_table, _add_page_break,
)

STATUS_COLOR = {"STRONG": None, "ADEQUATE": None, "BREACH": RED_TEXT}   # BREACH shown in red; others default black


def _pct(value: Optional[float]) -> str:
    return "" if value is None else f"{value * 100:.2f}%"


def _money(value: Optional[float]) -> str:
    return "" if value is None else f"GHS {value:,.2f}"


def _build_cover(doc: Document, client_name: str, valuation_date: str, report_date: str) -> None:
    for _ in range(4):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = 1  # WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("GIRBC Solvency Certificate")
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = STALLION_NAVY

    subtitle = doc.add_paragraph()
    subtitle.alignment = 1
    run = subtitle.add_run("Risk-Based Capital Adequacy — National Insurance Commission of Ghana")
    run.font.size = Pt(13)
    run.font.italic = True

    doc.add_paragraph()
    company = doc.add_paragraph()
    company.alignment = 1
    run = company.add_run(client_name)
    run.font.size = Pt(18)
    run.font.bold = True

    period = doc.add_paragraph()
    period.alignment = 1
    run = period.add_run(f"Valuation Date: {valuation_date}")
    run.font.size = Pt(12)

    date_p = doc.add_paragraph()
    date_p.alignment = 1
    run = date_p.add_run(f"Certificate Date: {report_date}")
    run.font.size = Pt(10)


def _build_summary(doc: Document, sol: SolvencyResult, legacy: Optional[LegacySolvencyResult]) -> None:
    _add_page_break(doc)
    _add_heading(doc, "1. Capital Adequacy Summary")

    car_p = doc.add_paragraph()
    run = car_p.add_run(f"GIRBC Capital Adequacy Ratio (CAR): {_pct(sol.car)}")
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = STATUS_COLOR.get(sol.status) or STALLION_NAVY

    status_p = doc.add_paragraph()
    run = status_p.add_run(f"Status: {sol.status}")
    run.font.size = Pt(12)
    run.font.bold = True
    if STATUS_COLOR.get(sol.status):
        run.font.color.rgb = STATUS_COLOR[sol.status]

    if legacy is not None:
        legacy_p = doc.add_paragraph()
        run = legacy_p.add_run(f"Legacy margin test (transition basis): {_pct(legacy.legacy_car)} — {legacy.status}")
        run.font.size = Pt(11)
        run.font.italic = True

    doc.add_paragraph()
    _add_kv_table(doc, [
        ("Total Qualifying Capital Resources", _money(sol.total_qcr)),
        ("Overall Risk Charge",                 _money(sol.overall_risk_charge)),
        ("Minimum Capital Requirement (MCR)",       _money(sol.mcr)),
        ("Prescribed Capital Requirement (PCR)",       _money(sol.pcr)),
        ("Surplus / (Deficit) vs MCR",                    _money(sol.total_qcr - sol.mcr)),
        ("Capital Composition Valid (Tier 1 >= 50% QCR)",    "Yes" if sol.capital_composition_valid else "No"),
    ])


def _build_risk_breakdown(doc: Document, sol: SolvencyResult) -> None:
    _add_heading(doc, "2. Risk Module Breakdown")
    naive_sum = (
        sol.insurance_risk.insurance_risk_before_cat + sol.insurance_risk.catastrophe_charge
        + sol.market_risk.total_market_risk_scr + sol.credit_risk.total_credit_risk_scr
    )
    diversification_benefit = naive_sum - sol.overall_risk_charge_before_operational
    _add_data_table(doc, ["Risk Module", "Charge (GHS)"], [
        ["Insurance Risk (excl. Catastrophe)", f"{sol.insurance_risk.insurance_risk_before_cat:,.2f}"],
        ["Catastrophe Risk", f"{sol.insurance_risk.catastrophe_charge:,.2f}"],
        ["Market Risk", f"{sol.market_risk.total_market_risk_scr:,.2f}"],
        ["Credit Risk", f"{sol.credit_risk.total_credit_risk_scr:,.2f}"],
        ["Subtotal (before diversification)", f"{naive_sum:,.2f}"],
        ["Diversification Benefit", f"({diversification_benefit:,.2f})"],
        ["Subtotal after diversification", f"{sol.overall_risk_charge_before_operational:,.2f}"],
        ["Operational Risk (no diversification credit)", f"{sol.operational_risk.total_operational_risk_scr:,.2f}"],
        ["Total Overall Risk Charge", f"{sol.overall_risk_charge:,.2f}"],
    ], total_row_indices={8})


def _build_capital_composition(doc: Document, sol: SolvencyResult) -> None:
    _add_heading(doc, "3. Capital Composition")
    cap = sol.capital_resources
    _add_data_table(doc, ["Component", "Amount (GHS)"], [
        ["Tier 1 Unlimited (gross)", f"{cap.tier1_unlimited:,.2f}"],
        ["Tier 1 Unlimited deductions", f"({cap.tier1_unlimited_deductions:,.2f})"],
        ["Net Tier 1 Unlimited", f"{cap.net_tier1_unlimited:,.2f}"],
        ["Tier 1 Limited (eligible, after 25% cap)", f"{cap.eligible_tier1_limited:,.2f}"],
        ["Tier 2 (eligible, after 25% cap)", f"{cap.eligible_tier2:,.2f}"],
        ["Total Qualifying Capital Resources", f"{cap.total_qcr:,.2f}"],
    ], total_row_indices={5})
    if cap.total_qcr > 0:
        tier1_pct = cap.net_tier1_unlimited / cap.total_qcr
        _add_body(doc, f"Tier 1 Unlimited represents {tier1_pct*100:.1f}% of total QCR "
                        f"(minimum 50% required — {'PASS' if tier1_pct >= 0.5 else 'BREACH'}).", bold=True)


def _build_stress_tests(doc: Document, stress_results: List[StressTestResult]) -> None:
    _add_heading(doc, "4. Stress Test Results")
    _add_body(doc, "Results of the 5 prescribed NIC stress scenarios applied to the base solvency position.")
    rows = []
    for r in stress_results:
        rows.append([
            r.scenario_name, _pct(r.new_car), r.status, "PASS" if r.passed else "FAIL",
            f"{r.car_change*100:+.2f}pp",
        ])
    _add_data_table(doc, ["Scenario", "Stressed CAR", "Status", "Result", "Change vs Base"], rows)


def _build_opinion(doc: Document, client_name: str, valuation_date: str, sol: SolvencyResult) -> None:
    _add_heading(doc, "5. Actuarial Opinion")
    solvent_text = (
        f"the Company remains above the Minimum Capital Requirement" if sol.car >= 1.0
        else f"the Company falls BELOW the Minimum Capital Requirement and is in BREACH of the "
             f"regulatory minimum solvency standard"
    )
    _add_body(doc,
        f"Based on the GIRBC risk-based capital calculation performed as at {valuation_date}, "
        f"{client_name}'s Qualifying Capital Resources of GHS {sol.total_qcr:,.2f} against an Overall "
        f"Risk Charge of GHS {sol.overall_risk_charge:,.2f} produce a Capital Adequacy Ratio of "
        f"{_pct(sol.car)}, resulting in a solvency status of {sol.status}. In my opinion, {solvent_text} "
        f"as at the valuation date.", size=11)
    if sol.status != "STRONG":
        _add_body(doc,
            "The Company's CAR is below the 150% Prescribed Capital Requirement (PCR) supervisory "
            "benchmark and should be monitored, with capital management actions considered where "
            "appropriate.", size=11, italic=True)


def _build_signature(doc: Document, appointed_actuary: str, qualifications: str, consulting_firm: str, report_date: str) -> None:
    doc.add_paragraph()
    doc.add_paragraph()
    sig_line = doc.add_paragraph()
    sig_line.add_run("_" * 45).font.size = Pt(11)
    for label, value in [
        ("Signed:", ""), ("Name:", appointed_actuary), ("Qualification:", qualifications),
        ("Firm:", consulting_firm), ("Date:", report_date),
    ]:
        p = doc.add_paragraph()
        run = p.add_run(f"{label}  ")
        run.font.bold = True
        run.font.size = Pt(11)
        run2 = p.add_run(str(value))
        run2.font.size = Pt(11)


def generate_girbc_certificate_from_results(
    client_name:          str,
    valuation_date:          str,
    sol:                        SolvencyResult,
    legacy_result:                 Optional[LegacySolvencyResult],
    stress_results:                    List[StressTestResult],
    appointed_actuary:                    str = "Charles Osei-Akoto",
    qualifications:                          str = "Fellow, Institute and Faculty of Actuaries (FIA)",
    consulting_firm:                            str = "Stallion Consultants Ltd",
    output_path:                                   Optional[str] = None,
) -> str:
    """Build the certificate from already-computed results (avoids a second engine run when the caller already has them)."""
    report_date = datetime.now().strftime("%d %B %Y")
    doc = Document()

    _build_cover(doc, client_name, valuation_date, report_date)
    _build_summary(doc, sol, legacy_result)
    _build_risk_breakdown(doc, sol)
    _build_capital_composition(doc, sol)
    _build_stress_tests(doc, stress_results)
    _build_opinion(doc, client_name, valuation_date, sol)
    _build_signature(doc, appointed_actuary, qualifications, consulting_firm, report_date)

    if output_path is None:
        os.makedirs(GENERATED_DIR, exist_ok=True)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        client_slug = client_name.replace(" ", "_").replace("/", "-")
        output_path = os.path.join(GENERATED_DIR, f"GIRBC_Certificate_{client_slug}_{valuation_date}_{timestamp}.docx")

    doc.save(output_path)
    return output_path


def generate_girbc_certificate(
    client_id:             str = "pic",
    valuation_date:            str = "FY2025",
    appointed_actuary:            str = "Charles Osei-Akoto",
    qualifications:                  str = "Fellow, Institute and Faculty of Actuaries (FIA)",
    consulting_firm:                    str = "Stallion Consultants Ltd",
    company_name_override:                 Optional[str] = None,
    output_path:                               Optional[str] = None,
) -> str:
    """
    Self-contained entry point: loads client_id's real GIRBC/legacy
    workbooks, runs the full engine/rbc/ pipeline, and writes the
    certificate. See generate_girbc_certificate_from_results() to reuse
    already-computed results instead.
    """
    client = load_client(client_id)
    display_name = company_name_override or client.name

    data = load_rbc_solvency_data(client_id)
    ins = calculate_insurance_risk(data["insurance_risk"], data["net_non_life_insurance_revenue"])
    mkt = calculate_market_risk(data["market_risk"])
    cred = calculate_credit_risk(data["credit_risk"])
    op = calculate_operational_risk(data["operational_risk"])
    sol = calculate_solvency(ins, mkt, cred, op, data["capital_resources"])

    legacy_result = calculate_legacy_solvency(data["legacy_inputs"]) if data["legacy_inputs"] is not None else None

    stress_results = run_stress_tests(
        data["insurance_risk"], data["net_non_life_insurance_revenue"],
        data["market_risk"], data["credit_risk"], data["operational_risk"], data["capital_resources"],
    )

    return generate_girbc_certificate_from_results(
        client_name=display_name, valuation_date=valuation_date, sol=sol, legacy_result=legacy_result,
        stress_results=stress_results, appointed_actuary=appointed_actuary, qualifications=qualifications,
        consulting_firm=consulting_firm, output_path=output_path,
    )
